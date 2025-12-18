from flask import Flask, request, jsonify
import requests
from io import BytesIO
from docx import Document

app = Flask(__name__)


@app.route("/extract-job-description", methods=["POST"])
def extract_job_description():
    data = request.get_json(silent=True)

    if not data:
        return jsonify({"error": "Invalid JSON body"}), 400

    file_url = data.get("file_url")
    api_key = data.get("api_key")

    if not file_url or not api_key:
        return jsonify({"error": "file_url and api_key are required"}), 400

    try:
        response = requests.get(
            file_url,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Accept": "*/*",
            },
            timeout=20,
        )

        if response.status_code != 200:
            return jsonify({"error": "Failed to download DOCX file"}), 400

        doc = Document(BytesIO(response.content))
        text_parts = []

        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())

        # Extract tables (CRITICAL for Recruit CRM)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        extracted_text = "\n".join(text_parts)

        if not extracted_text.strip():
            return jsonify({"error": "No text found in DOCX"}), 422

        return jsonify({"extracted_text": extracted_text})

    except Exception as e:
        return jsonify({"error": str(e)}), 500
